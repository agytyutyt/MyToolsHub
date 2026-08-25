"""共享文档 —— JZToolsHub 后端插件路由。

功能：支持多用户共同编辑 Word（.docx）与 Excel（.xlsx）文档：
- 文档以 JSON 文件保存在插件目录 data/ 下，每个文档一个文件；
- 内容带版本号（乐观锁）：保存时携带 base_version，版本冲突返回 409，
  由前端提示用户「放弃本地修改」或「覆盖保存」，避免互相静默覆盖丢内容；
- 在线协作：客户端定时心跳上报「正在编辑」，后端维护在线用户列表；
- 导入 / 导出：Word 用 python-docx，Excel 用 openpyxl / xlrd，
  依赖缺失时对应接口返回明确安装提示（不影响纯编辑）。

内容格式（JSON）：
- word :  {"blocks": [{"type": "p|h1..h6|ul|ol",
             "runs": [{"t": "文本", "b": bool, "i": bool, "u": bool}],
             "items": [{"runs": [...]}]}]}
- excel:  {"rows": [[单元格, ...], ...]}   空单元格统一存空字符串
"""

import io
import json
import os
import re
import threading
import time
import uuid
from datetime import datetime

from flask import jsonify, request, send_file

try:
    from jztools_admin.routes import get_session_user as _get_session_user
except Exception:  # admin 插件缺失时兜底（理论上不会发生）
    _get_session_user = None

API_PREFIX = "/api/shared-docs"

_BACKEND_DIR = os.path.dirname(os.path.abspath(__file__))
_DATA_DIR = os.path.join(_BACKEND_DIR, "data")

_DOC_ID_RE = re.compile(r"^[A-Za-z0-9_-]+$")

# 全局写锁：文档文件读写、版本递增、在线用户维护共用一把锁
_LOCK = threading.RLock()

# 在线用户（内存态）：doc_id -> {client_id: {"user": ..., "last_seen": ts}}
PRESENCE = {}
PRESENCE_TTL = 20  # 秒；超过该时长未心跳视为离线

MAX_DOC_NAME = 100
MAX_WORD_BLOCKS = 5000
MAX_EXCEL_ROWS = 1000
MAX_EXCEL_COLS = 100
MAX_HISTORY = 100


def _viewer():
    """当前登录用户上下文；未登录返回 None。"""
    if _get_session_user is None:
        return None
    try:
        return _get_session_user()
    except Exception:
        return None


def _can_view(user, doc):
    """按可见性规则判定 user 是否可见 doc（3.2-③）。

    - unit 级：同单位可见；department 级：同部门可见；private：仅创建者；
    - 超级管理员可见全部。
    """
    if user is None:
        return False
    if user.get("super_admin"):
        return True
    scope = doc.get("scope") or {}
    level = scope.get("level", "private")
    if level == "unit":
        return scope.get("unit_id") == user.get("unit_id")
    if level == "department":
        return scope.get("department_id") == user.get("department_id")
    return scope.get("owner") == user.get("username")


def _can_manage(user, doc):
    """重命名/删除/改挂靠：仅创建者或超级管理员。"""
    if user is None:
        return False
    if user.get("super_admin"):
        return True
    return (doc.get("created_by") or "") == user.get("username")

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except Exception:  # 依赖缺失时导出/导入接口给出安装提示
    Document = None
    qn = None
    Pt = None
    DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import Font
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except Exception:
    openpyxl = None
    Font = None
    get_column_letter = None
    OPENPYXL_AVAILABLE = False

try:
    import xlrd
    XLRD_AVAILABLE = True
except Exception:
    xlrd = None
    XLRD_AVAILABLE = False


# ===================== 文档存储 =====================

def _now():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _doc_path(doc_id):
    return os.path.join(_DATA_DIR, f"{doc_id}.json")


def _check_doc_id(doc_id):
    return bool(_DOC_ID_RE.match(doc_id or ""))


def load_doc(doc_id):
    if not _check_doc_id(doc_id):
        return None
    path = _doc_path(doc_id)
    if not os.path.isfile(path):
        return None
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def save_doc(doc):
    """原子写：先写临时文件再替换，避免半截 JSON 落盘。"""
    os.makedirs(_DATA_DIR, exist_ok=True)
    path = _doc_path(doc["id"])
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(doc, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)


def list_docs():
    """列出全部文档，按更新时间倒序。"""
    os.makedirs(_DATA_DIR, exist_ok=True)
    docs = []
    for name in os.listdir(_DATA_DIR):
        if not name.endswith(".json") or name.endswith(".tmp.json"):
            continue
        doc_id = name[:-5]
        if not _check_doc_id(doc_id):
            continue
        try:
            with open(os.path.join(_DATA_DIR, name), "r", encoding="utf-8") as f:
                docs.append(json.load(f))
        except Exception:
            continue
    docs.sort(key=lambda d: d.get("updated_at", ""), reverse=True)
    return docs


def migrate_doc_scope():
    """给没有 scope 的历史文档补默认归属：私有，创建者 admin。

    历史文档无法追溯创建人与受众，默认划给超管 admin 最安全；
    现场需要老文档全员共享时由管理员在前端改成「单位」层级即可。
    """
    changed = False
    for doc in list_docs():
        if doc.get("scope"):
            continue
        doc["created_by"] = doc.get("created_by") or "admin"
        doc["created_by_name"] = doc.get("created_by_name") or "系统管理员"
        doc["scope"] = {
            "level": "private",
            "owner": "admin",
            "owner_name": "系统管理员",
            "unit_id": "",
            "department_id": "",
        }
        save_doc(doc)
        changed = True
    return changed


def _public_doc(doc, with_content=False, with_presence=False):
    payload = {
        "id": doc["id"],
        "name": doc["name"],
        "type": doc["type"],
        "version": doc.get("version", 0),
        "created_at": doc.get("created_at", ""),
        "updated_at": doc.get("updated_at", ""),
        "updated_by": doc.get("updated_by", ""),
        "created_by": doc.get("created_by", ""),
        "created_by_name": doc.get("created_by_name", ""),
        "scope": doc.get("scope") or {},
    }
    if with_content:
        payload["content"] = doc.get("content", {})
        payload["history"] = doc.get("history", [])
    if with_presence:
        payload["users"] = presence_users(doc["id"])
    return payload


def _normalize_runs(runs):
    """规范化 runs：只保留 {t, b, i, u}，全部转为布尔。"""
    out = []
    if not isinstance(runs, list):
        return out
    for r in runs:
        if not isinstance(r, dict):
            continue
        t = r.get("t")
        if t is None:
            t = ""
        if not isinstance(t, str):
            t = str(t)
        if not t:
            continue
        out.append({
            "t": t,
            "b": bool(r.get("b")),
            "i": bool(r.get("i")),
            "u": bool(r.get("u")),
        })
    return out


def validate_content(doc_type, content):
    """校验并规范化内容，非法时抛 ValueError。"""
    if not isinstance(content, dict):
        raise ValueError("内容格式非法")
    if doc_type == "word":
        blocks = content.get("blocks")
        if not isinstance(blocks, list) or len(blocks) > MAX_WORD_BLOCKS:
            raise ValueError(f"Word 内容格式非法或超过 {MAX_WORD_BLOCKS} 个段落上限")
        norm = []
        for b in blocks:
            if not isinstance(b, dict):
                continue
            t = b.get("type", "p")
            if t not in ("p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol"):
                t = "p"
            if t in ("ul", "ol"):
                items = []
                for it in (b.get("items") or []):
                    if isinstance(it, dict):
                        items.append({"runs": _normalize_runs(it.get("runs"))})
                norm.append({"type": t, "items": items})
            else:
                norm.append({"type": t, "runs": _normalize_runs(b.get("runs"))})
        return {"blocks": norm}
    if doc_type == "excel":
        rows = content.get("rows")
        if not isinstance(rows, list):
            raise ValueError("Excel 内容格式非法")
        if len(rows) > MAX_EXCEL_ROWS:
            raise ValueError(f"Excel 行数超过上限 {MAX_EXCEL_ROWS}")
        norm = []
        for row in rows:
            if not isinstance(row, list):
                raise ValueError("Excel 行格式非法")
            if len(row) > MAX_EXCEL_COLS:
                raise ValueError(f"Excel 列数超过上限 {MAX_EXCEL_COLS}")
            norm.append(["" if v is None else v for v in row])
        # 列宽（像素）：0 表示自动，与列数对齐
        colWidths = content.get("colWidths")
        norm_widths = []
        if isinstance(colWidths, list):
            for w in colWidths:
                try:
                    w = float(w)
                except (TypeError, ValueError):
                    w = 0
                norm_widths.append(round(w, 1) if 20 <= w <= 600 else 0)
        return {"rows": norm, "colWidths": norm_widths}
    raise ValueError("未知文档类型")


# ===================== 在线用户 =====================

def presence_users(doc_id):
    now = time.time()
    info = PRESENCE.get(doc_id, {})
    users = []
    for cid, val in list(info.items()):
        if now - val["last_seen"] > PRESENCE_TTL:
            info.pop(cid, None)
        else:
            users.append(val["user"])
    return sorted(set(users))


# ===================== Word 导入 / 导出 =====================

def _set_east_asia_font(style, name):
    """为 docx 样式设置中文字体（w:eastAsia），避免中文回退字体异常。"""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _add_runs(paragraph, runs):
    for r in runs:
        run = paragraph.add_run(r.get("t") or "")
        run.bold = bool(r.get("b"))
        run.italic = bool(r.get("i"))
        run.underline = bool(r.get("u"))


def export_word(doc):
    """blocks → .docx 字节流。返回 (BytesIO, 文档名)。"""
    if not DOCX_AVAILABLE:
        raise RuntimeError("后端缺少 python-docx，无法导出 .docx，请执行：pip install python-docx")
    d = Document()
    style = d.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    _set_east_asia_font(style, "微软雅黑")

    for b in doc["content"].get("blocks", []):
        btype = b.get("type", "p")
        if btype in ("ul", "ol"):
            list_style = "List Bullet" if btype == "ul" else "List Number"
            for item in b.get("items") or []:
                p = d.add_paragraph(style=list_style)
                _add_runs(p, item.get("runs") or [])
            continue
        if btype in ("h1", "h2", "h3", "h4", "h5", "h6"):
            p = d.add_heading("", level=int(btype[1:]))
        else:
            p = d.add_paragraph()
        _add_runs(p, b.get("runs") or [])

    buf = io.BytesIO()
    d.save(buf)
    return buf, doc["name"]


def import_word(file_storage):
    """.docx → blocks。返回 {"blocks": [...]}。"""
    if not DOCX_AVAILABLE:
        raise RuntimeError("后端缺少 python-docx，无法导入 .docx，请执行：pip install python-docx")
    file_storage.seek(0)
    d = Document(file_storage)
    blocks = []
    for para in d.paragraphs:
        text = para.text or ""
        style_name = (para.style.name or "").lower().strip()
        runs = []
        for r in para.runs:
            if r.text:
                runs.append({"t": r.text, "b": r.bold, "i": r.italic, "u": r.underline})
        if not text and not runs:
            continue
        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip() or "1")
            except ValueError:
                level = 1
            blocks.append({"type": f"h{min(max(level, 1), 6)}", "runs": runs or [{"t": text}]})
        elif "list" in style_name:
            btype = "ol" if "number" in style_name else "ul"
            if blocks and blocks[-1].get("type") == btype:
                blocks[-1]["items"].append({"runs": runs or [{"t": text}]})
            else:
                blocks.append({"type": btype, "items": [{"runs": runs or [{"t": text}]}]})
        else:
            blocks.append({"type": "p", "runs": runs or [{"t": text}]})
    return {"blocks": blocks}


# ===================== Excel 导入 / 导出 =====================

def _cell_value(v):
    """把导入读到的单元格值统一为可 JSON 序列化的形式。"""
    if v is None:
        return ""
    if isinstance(v, datetime):
        return v.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(v, (bool, int, float)):
        return v
    return str(v)


def export_excel(doc):
    """rows → .xlsx 字节流。返回 (BytesIO, 文档名)。"""
    if not OPENPYXL_AVAILABLE:
        raise RuntimeError("后端缺少 openpyxl，无法导出 .xlsx，请执行：pip install openpyxl")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    rows = doc["content"].get("rows", [])
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            if val == "" or val is None:
                continue
            v = val
            if isinstance(v, str):
                # 纯数字字符串自动转成数值，导出后可直接参与公式计算
                try:
                    v = float(v)
                except (TypeError, ValueError):
                    pass
            ws.cell(row=i + 1, column=j + 1, value=v)
    if rows:
        for j in range(len(rows[0])):
            ws.cell(row=1, column=j + 1).font = Font(bold=True)
        ws.freeze_panes = "A2"
    # 应用在线编辑时调整过的列宽（0 表示自动，不写）
    for j, w in enumerate(doc["content"].get("colWidths") or []):
        if w:
            # openpyxl 列宽单位 ≈ 字符数（约为像素 / 7）
            ws.column_dimensions[get_column_letter(j + 1)].width = max(5, round(float(w) / 7, 1))
    buf = io.BytesIO()
    wb.save(buf)
    return buf, doc["name"]


def import_excel(file_storage, filename):
    """.xlsx / .xls → rows。返回 {"rows": [...]}。"""
    ext = (filename or "").lower().rsplit(".", 1)[-1]
    if ext not in ("xlsx", "xls"):
        raise ValueError("仅支持 .xlsx / .xls 文件")
    file_storage.seek(0)
    if ext == "xlsx":
        if not OPENPYXL_AVAILABLE:
            raise RuntimeError("后端缺少 openpyxl，无法导入 .xlsx，请执行：pip install openpyxl")
        wb = openpyxl.load_workbook(file_storage, data_only=True)
        try:
            ws = wb.active
            rows = [[_cell_value(v) for v in r] for r in ws.iter_rows(values_only=True)]
        finally:
            wb.close()
        return {"rows": rows}
    # .xls 走 xlrd（不支持直接从流读取，落临时文件）
    if not XLRD_AVAILABLE:
        raise RuntimeError("后端缺少 xlrd，无法导入 .xls，请执行：pip install xlrd")
    import tempfile
    tmp = tempfile.NamedTemporaryFile(suffix=".xls", delete=False)
    try:
        tmp.write(file_storage.read())
        tmp.close()
        book = xlrd.open_workbook(tmp.name)
        sheet = book.sheet_by_index(0)
        rows = []
        for r in range(sheet.nrows):
            row = []
            for c in range(sheet.ncols):
                cell = sheet.cell(r, c)
                if cell.ctype == xlrd.XL_CELL_DATE:
                    row.append(xlrd.xldate_as_datetime(cell.value, book.datemode).strftime("%Y-%m-%d %H:%M:%S"))
                elif cell.ctype == xlrd.XL_CELL_EMPTY:
                    row.append("")
                else:
                    row.append(_cell_value(cell.value))
            rows.append(row)
        return {"rows": rows}
    finally:
        os.remove(tmp.name)


# ===================== 路由注册 =====================

def register(app) -> None:
    """插件入口：由 JZToolsHub 主应用在启动时调用。"""
    # 存量文档归属迁移（一次性、幂等）：给历史文档补默认归属
    with _LOCK:
        migrate_doc_scope()

    @app.get(f"{API_PREFIX}/status")
    def sd_status():
        return jsonify({
            "ok": True,
            "python_docx": DOCX_AVAILABLE,
            "openpyxl": OPENPYXL_AVAILABLE,
            "xlrd": XLRD_AVAILABLE,
        })

    @app.get(f"{API_PREFIX}/documents")
    def sd_list():
        user = _viewer()
        if user is None:
            return jsonify({"ok": False, "detail": "未登录或登录已过期"}), 401
        docs = [d for d in list_docs() if _can_view(user, d)]
        return jsonify({"ok": True, "documents": [_public_doc(d) for d in docs]})

    @app.post(f"{API_PREFIX}/documents")
    def sd_create():
        user = _viewer()
        if user is None:
            return jsonify({"ok": False, "detail": "未登录或登录已过期"}), 401
        data = request.get_json(silent=True) or {}
        name = (data.get("name") or "").strip()
        doc_type = data.get("type")
        if not name:
            return jsonify({"ok": False, "detail": "请输入文档名称"}), 400
        if len(name) > MAX_DOC_NAME:
            return jsonify({"ok": False, "detail": f"文档名称不能超过 {MAX_DOC_NAME} 个字符"}), 400
        if doc_type not in ("word", "excel"):
            return jsonify({"ok": False, "detail": "文档类型必须为 word 或 excel"}), 400
        # 挂靠层级只能三选一，组织信息一律取自 session（不信前端）
        level = (data.get("level") or "private").strip()
        if level not in ("unit", "department", "private"):
            return jsonify({"ok": False, "detail": "挂靠层级必须为 unit / department / private"}), 400
        doc = {
            "id": uuid.uuid4().hex[:12],
            "name": name,
            "type": doc_type,
            "version": 0,
            "content": {"blocks": []} if doc_type == "word" else {"rows": [[""]]},
            "created_at": _now(),
            "updated_at": _now(),
            "updated_by": user["name"],
            "history": [],
            "created_by": user["username"],
            "created_by_name": user["name"],
            "scope": {
                "level": level,
                "owner": user["username"],
                "owner_name": user["name"],
                "unit_id": user.get("unit_id", ""),
                "department_id": user.get("department_id", ""),
            },
        }
        save_doc(doc)
        return jsonify({"ok": True, "document": _public_doc(doc, with_content=True)})

    @app.get(f"{API_PREFIX}/documents/<doc_id>")
    def sd_get(doc_id):
        user = _viewer()
        if user is None:
            return jsonify({"ok": False, "detail": "未登录或登录已过期"}), 401
        doc = load_doc(doc_id)
        if not doc or not _can_view(user, doc):
            return jsonify({"ok": False, "detail": "文档不存在或无权访问"}), 404
        return jsonify({"ok": True, "document": _public_doc(doc, with_content=True, with_presence=True)})

    @app.post(f"{API_PREFIX}/documents/<doc_id>/content")
    def sd_save_content(doc_id):
        """保存内容（乐观锁）：base_version 与当前版本不一致时返回 409。"""
        user = _viewer()
        if user is None:
            return jsonify({"ok": False, "detail": "未登录或登录已过期"}), 401
        data = request.get_json(silent=True) or {}
        base_version = data.get("base_version")
        with _LOCK:
            doc = load_doc(doc_id)
            if not doc or not _can_view(user, doc):
                return jsonify({"ok": False, "detail": "文档不存在或无权访问"}), 404
            if not isinstance(base_version, int) or base_version != doc.get("version", 0):
                return jsonify({
                    "ok": False,
                    "detail": "文档已被其他用户更新，请选择「丢弃本地修改」或「覆盖保存」",
                    "document": _public_doc(doc, with_content=True),
                }), 409
            try:
                content = validate_content(doc["type"], data.get("content"))
            except ValueError as e:
                return jsonify({"ok": False, "detail": str(e)}), 400
            doc["version"] = doc.get("version", 0) + 1
            doc["content"] = content
            doc["updated_at"] = _now()
            doc["updated_by"] = user["name"]
            history = doc.setdefault("history", [])
            history.append({"version": doc["version"], "time": doc["updated_at"], "by": user["name"]})
            if len(history) > MAX_HISTORY:
                doc["history"] = history[-MAX_HISTORY:]
            save_doc(doc)
        return jsonify({"ok": True, "document": _public_doc(doc)})

    @app.post(f"{API_PREFIX}/documents/<doc_id>/presence")
    def sd_presence(doc_id):
        """在线心跳：前端定时调用，返回当前文档在线用户。

        在线用户身份只认 session（忽略前端自报昵称），防冒名顶替。
        """
        user = _viewer()
        if user is None:
            return jsonify({"ok": False, "detail": "未登录或登录已过期"}), 401
        data = request.get_json(silent=True) or {}
        client_id = (data.get("client_id") or "unknown")[:64]
        with _LOCK:
            doc = load_doc(doc_id)
            if not doc or not _can_view(user, doc):
                return jsonify({"ok": False, "detail": "文档不存在或无权访问"}), 404
            now = time.time()
            PRESENCE.setdefault(doc_id, {})[client_id] = {
                "user": user["name"], "username": user["username"], "last_seen": now,
            }
            users = presence_users(doc_id)
        return jsonify({"ok": True, "users": users})

    @app.post(f"{API_PREFIX}/documents/<doc_id>/rename")
    def sd_rename(doc_id):
        user = _viewer()
        if user is None:
            return jsonify({"ok": False, "detail": "未登录或登录已过期"}), 401
        data = request.get_json(silent=True) or {}
        name = (data.get("name") or "").strip()
        if not name:
            return jsonify({"ok": False, "detail": "请输入文档名称"}), 400
        if len(name) > MAX_DOC_NAME:
            return jsonify({"ok": False, "detail": f"文档名称不能超过 {MAX_DOC_NAME} 个字符"}), 400
        with _LOCK:
            doc = load_doc(doc_id)
            if not doc or not _can_view(user, doc):
                return jsonify({"ok": False, "detail": "文档不存在或无权访问"}), 404
            if not _can_manage(user, doc):
                return jsonify({"ok": False, "detail": "仅创建者或管理员可执行此操作"}), 403
            doc["name"] = name
            save_doc(doc)
        return jsonify({"ok": True, "document": _public_doc(doc)})

    @app.delete(f"{API_PREFIX}/documents/<doc_id>")
    def sd_delete(doc_id):
        user = _viewer()
        if user is None:
            return jsonify({"ok": False, "detail": "未登录或登录已过期"}), 401
        with _LOCK:
            doc = load_doc(doc_id)
            if not doc or not _can_view(user, doc):
                return jsonify({"ok": False, "detail": "文档不存在或无权访问"}), 404
            if not _can_manage(user, doc):
                return jsonify({"ok": False, "detail": "仅创建者或管理员可执行此操作"}), 403
            try:
                os.remove(_doc_path(doc_id))
            except OSError:
                return jsonify({"ok": False, "detail": "删除失败"}), 500
        return jsonify({"ok": True})

    @app.post(f"{API_PREFIX}/documents/<doc_id>/scope")
    def sd_change_scope(doc_id):
        """调整文档挂靠层级（unit / department / private）：仅创建者或超级管理员。

        只改 level，保留原始 owner/unit/dept 组织信息（语义是「调整共享范围」，
        而不是「把文档搬到别的单位」）。
        """
        body = request.get_json(silent=True) or {}
        level = (body.get("level") or "").strip()
        if level not in ("unit", "department", "private"):
            return jsonify({"ok": False, "detail": "挂靠层级不合法"}), 400
        with _LOCK:
            user = _viewer()
            if user is None:
                return jsonify({"ok": False, "detail": "未登录或登录已过期"}), 401
            doc = load_doc(doc_id)
            if not doc or not _can_view(user, doc):
                return jsonify({"ok": False, "detail": "文档不存在或无权访问"}), 404
            if not _can_manage(user, doc):
                return jsonify({"ok": False, "detail": "仅创建者或管理员可调整挂靠"}), 403
            doc.setdefault("scope", {})["level"] = level
            save_doc(doc)
        return jsonify({"ok": True, "document": _public_doc(doc)})

    @app.get(f"{API_PREFIX}/documents/<doc_id>/export")
    def sd_export(doc_id):
        """导出为真实 Office 文件（.docx / .xlsx）。"""
        user = _viewer()
        if user is None:
            return jsonify({"ok": False, "detail": "未登录或登录已过期"}), 401
        doc = load_doc(doc_id)
        if not doc or not _can_view(user, doc):
            return jsonify({"ok": False, "detail": "文档不存在或无权访问"}), 404
        try:
            if doc["type"] == "word":
                buf, name = export_word(doc)
                mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ext = "docx"
            else:
                buf, name = export_excel(doc)
                mimetype = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                ext = "xlsx"
        except RuntimeError as e:
            return jsonify({"ok": False, "detail": str(e)}), 400
        buf.seek(0)
        return send_file(buf, as_attachment=True,
                         download_name=f"{name}.{ext}", mimetype=mimetype)

    @app.post(f"{API_PREFIX}/documents/<doc_id>/import")
    def sd_import(doc_id):
        """从 .docx / .xlsx / .xls 导入内容，覆盖当前文档。

        可见即可编辑：能看到的单位/部门/私人文档（本人）均可导入覆盖。
        """
        user = _viewer()
        if user is None:
            return jsonify({"ok": False, "detail": "未登录或登录已过期"}), 401
        f = request.files.get("file")
        if not f or not f.filename:
            return jsonify({"ok": False, "detail": "请选择文件"}), 400
        ext = f.filename.lower().rsplit(".", 1)[-1]
        with _LOCK:
            doc = load_doc(doc_id)
            if not doc or not _can_view(user, doc):
                return jsonify({"ok": False, "detail": "文档不存在或无权访问"}), 404
            if doc["type"] == "word" and ext != "docx":
                return jsonify({"ok": False, "detail": "Word 文档仅支持导入 .docx"}), 400
            if doc["type"] == "excel" and ext not in ("xlsx", "xls"):
                return jsonify({"ok": False, "detail": "Excel 文档仅支持导入 .xlsx / .xls"}), 400
            try:
                if doc["type"] == "word":
                    content = import_word(f)
                else:
                    content = import_excel(f, f.filename)
                content = validate_content(doc["type"], content)
            except (ValueError, RuntimeError) as e:
                return jsonify({"ok": False, "detail": str(e)}), 400
            doc["version"] = doc.get("version", 0) + 1
            doc["content"] = content
            doc["updated_at"] = _now()
            doc["updated_by"] = user["name"]
            history = doc.setdefault("history", [])
            history.append({"version": doc["version"], "time": doc["updated_at"], "by": doc["updated_by"]})
            if len(history) > MAX_HISTORY:
                doc["history"] = history[-MAX_HISTORY:]
            save_doc(doc)
        return jsonify({"ok": True, "document": _public_doc(doc, with_content=True)})
